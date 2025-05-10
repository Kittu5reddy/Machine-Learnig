

from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Image Feature Engineering: Techniques and Examples', 0)

# Adding an introduction section
doc.add_paragraph(
    "This document provides an overview of common image feature engineering techniques used in computer vision. "
    "These techniques include edge detection, texture analysis, and Histogram of Oriented Gradients (HOG). "
    "Each technique is followed by an example implementation."
)

# Adding Edge Detection section
doc.add_heading('1. Edge Detection', level=1)
doc.add_paragraph(
    "Edge detection is a technique used to identify the boundaries or edges of objects within an image. "
    "It is particularly useful for object detection, image segmentation, and boundary detection."
)
doc.add_paragraph(
    "Example using Canny Edge Detection (Python code):\n"
    "```python\n"
    "import cv2\n"
    "import matplotlib.pyplot as plt\n\n"
    "image = cv2.imread('image_path')\n"
    "gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)\n"
    "edges = cv2.Canny(gray, threshold1=100, threshold2=200)\n\n"
    "plt.imshow(edges, cmap='gray')\n"
    "plt.title('Canny Edge Detection')\n"
    "plt.show()\n"
    "```"
)

# Adding Texture Analysis section
doc.add_heading('2. Texture Analysis', level=1)
doc.add_paragraph(
    "Texture analysis helps in identifying patterns and repetitive structures in an image. "
    "It is important in applications like texture classification and face recognition."
)
doc.add_paragraph(
    "Example using Gray Level Co-occurrence Matrix (GLCM) (Python code):\n"
    "```python\n"
    "import numpy as np\n"
    "import skimage.feature as feature\n\n"
    "image = np.array([[0, 1, 2], [3, 4, 5], [6, 7, 8]])  # Example image\n"
    "glcm = feature.greycomatrix(image, distances=[1], angles=[0], symmetric=True, normed=True)\n"
    "contrast = feature.greycoprops(glcm, prop='contrast')\n"
    "print('GLCM Contrast:', contrast)\n"
    "```"
)

# Adding HOG (Histogram of Oriented Gradients) section
doc.add_heading('3. Histogram of Oriented Gradients (HOG)', level=1)
doc.add_paragraph(
    "HOG is a feature descriptor that counts occurrences of gradient orientation in localized portions of an image. "
    "It is particularly effective for human detection and object recognition."
)
doc.add_paragraph(
    "Example using HOG (Python code):\n"
    "```python\n"
    "from skimage.feature import hog\n"
    "from skimage import data, exposure\n\n"
    "image = data.astronaut()  # Sample image\n"
    "fd, hog_image = hog(image, pixels_per_cell=(16, 16), cells_per_block=(1, 1), visualize=True)\n\n"
    "hog_image_rescaled = exposure.rescale_intensity(hog_image, in_range=(0, 10))\n"
    "plt.imshow(hog_image_rescaled, cmap=plt.cm.gray)\n"
    "plt.title('Histogram of Oriented Gradients (HOG)')\n"
    "plt.show()\n"
    "```"
)

# Save the document
doc.save('Image_Feature_Engineering_Examples.docx')


